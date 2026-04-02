"""Convert README.md to a properly structured Word .docx"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.54)
section.top_margin  = section.bottom_margin = Cm(2.54)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_run_font(run, name_en='Calibri', name_cn='宋体', size=None, bold=False,
                 italic=False, color=None):
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    if size:  run.font.size = Pt(size)
    if bold:  run.font.bold = bold
    if italic: run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def heading(text, level):
    """Add a Word heading (Heading 1/2/3)."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def para(text='', bold_segments=None, style='Normal'):
    """Add a normal paragraph, supporting **bold** inline markers."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    if not text:
        return p
    # Split on **...**
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = (i % 2 == 1)
        set_run_font(run, size=11, bold=is_bold)
    return p

def code_block(text):
    """Add a grey-background code paragraph."""
    for line in text.split('\n'):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        # grey shading
        rPr = run._element.rPr
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        rPr.append(shd)

def add_table(headers, rows):
    """Add a styled table."""
    col_n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=col_n)
    t.style = 'Table Grid'
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        rPrh = run._element.get_or_add_rPr()
        rfh = rPrh.find(qn('w:rFonts'))
        if rfh is None:
            rfh = OxmlElement('w:rFonts'); rPrh.insert(0, rfh)
        rfh.set(qn('w:eastAsia'), '黑体')
        # header bg
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data[:col_n]):
            cell = row.cells[ci]
            # strip markdown bold
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', str(val))
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            cell.text = clean
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
            else:
                run = cell.paragraphs[0].add_run(clean)
            run.font.size = Pt(10)
            rPr2 = run._element.get_or_add_rPr()
            rf2 = rPr2.find(qn('w:rFonts'))
            if rf2 is None:
                rf2 = OxmlElement('w:rFonts'); rPr2.insert(0, rf2)
            rf2.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()  # spacing after table

def bullet(text, level=0):
    """Add a bullet list item."""
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    clean = re.sub(r'`(.+?)`', r'\1', clean)
    run = p.add_run(clean)
    set_run_font(run, size=11)

def checkbox_item(text):
    """Render [ ] or [x] as a bullet."""
    text = re.sub(r'^\[.\] ', '☐ ', text)
    bullet(text)

def divider():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Parse and render README.md
# ══════════════════════════════════════════════════════════════════════════════
with open('E:/Xgboost/README.md', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code = False
code_buf = []
in_table = False
table_headers = []
table_rows = []

def flush_table():
    global in_table, table_headers, table_rows
    if table_headers:
        add_table(table_headers, table_rows)
    in_table = False
    table_headers = []
    table_rows = []

def flush_code():
    global in_code, code_buf
    if code_buf:
        code_block('\n'.join(code_buf))
    in_code = False
    code_buf = []

while i < len(lines):
    raw = lines[i].rstrip('\n')
    stripped = raw.strip()

    # ── Code fence ──────────────────────────────────────────────────────────
    if stripped.startswith('```'):
        if not in_code:
            if in_table: flush_table()
            in_code = True
            code_buf = []
        else:
            flush_code()
        i += 1
        continue

    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    # ── Table ────────────────────────────────────────────────────────────────
    if stripped.startswith('|'):
        # separator row
        if re.match(r'^\|[-| :]+\|$', stripped):
            i += 1
            continue
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        if not in_table:
            in_table = True
            table_headers = cells
            table_rows = []
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            flush_table()

    # ── Headings ─────────────────────────────────────────────────────────────
    m = re.match(r'^(#{1,3})\s+(.*)', stripped)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        # skip TOC heading
        if text == '目录':
            # skip until next blank-line-separated section
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#'):
                i += 1
            continue
        heading(text, level)
        i += 1
        continue

    # ── Horizontal rule ───────────────────────────────────────────────────────
    if stripped in ('---', '***', '___'):
        divider()
        i += 1
        continue

    # ── Blockquote ────────────────────────────────────────────────────────────
    if stripped.startswith('> '):
        text = stripped[2:]
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(re.sub(r'`(.+?)`', r'\1', text))
        set_run_font(run, size=10.5, italic=True, color=(0x60, 0x60, 0x60))
        i += 1
        continue

    # ── Checkbox items ────────────────────────────────────────────────────────
    if re.match(r'^- \[.\]', stripped):
        checkbox_item(stripped[2:])
        i += 1
        continue

    # ── Bullet list ───────────────────────────────────────────────────────────
    if stripped.startswith('- ') or stripped.startswith('* '):
        checkbox_item(stripped[2:]) if stripped[2:4] in ('[x]', '[ ]') \
            else bullet(stripped[2:])
        i += 1
        continue

    # ── Numbered list ─────────────────────────────────────────────────────────
    if re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        # If it looks like a TOC entry, skip
        if re.search(r'\[.*?\]\(#.*?\)', text):
            i += 1
            continue
        bullet(text)
        i += 1
        continue

    # ── Empty line ────────────────────────────────────────────────────────────
    if not stripped:
        i += 1
        continue

    # ── Footer italic line ────────────────────────────────────────────────────
    if stripped.startswith('*') and stripped.endswith('*') and stripped.count('*') == 2:
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(stripped.strip('*'))
        set_run_font(run, size=9, italic=True, color=(0x80, 0x80, 0x80))
        i += 1
        continue

    # ── Normal paragraph ─────────────────────────────────────────────────────
    para(stripped)
    i += 1

# flush any open states
if in_table: flush_table()
if in_code:  flush_code()

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'E:/Xgboost/README.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
